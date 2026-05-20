import io
import json
import os
import re

import anthropic
from docx import Document
from docx.shared import Inches, Pt
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.orm import Session

import models
import schemas
from database import get_db

router = APIRouter(prefix="/deals", tags=["memo"])

_client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
_SONNET = "claude-sonnet-4-6"


def _latest_outputs_by_agent(agent_outputs: list) -> dict:
    seen: dict = {}
    for ao in sorted(agent_outputs, key=lambda x: x.created_at, reverse=True):
        if ao.agent_name not in seen:
            seen[ao.agent_name] = json.loads(ao.output_json)
    return seen


def _latest_feedback_by_agent(feedback_entries: list) -> dict:
    seen: dict = {}
    for fb in sorted(feedback_entries, key=lambda x: x.created_at, reverse=True):
        if fb.agent_name not in seen:
            seen[fb.agent_name] = fb.feedback_text
    return seen


def _parse_json(raw: str) -> dict:
    raw = raw.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            try:
                return json.loads(m.group())
            except json.JSONDecodeError:
                return {}
        return {}


def _bullet(items: list) -> str:
    return "\n".join(f"- {item}" for item in items) if items else "- N/A"


_RATING_LABELS = {
    "conviction": "Conviction & Commitment",
    "expertise": "Scientific / Domain Expertise",
    "execution": "Execution Track Record",
    "vision": "Strategic Clarity",
    "team": "Team Cohesion",
    "coachability": "Openness to Input",
}

_SCORE_LABELS = {1: "Unclear", 2: "Developing", 3: "Adequate", 4: "Strong", 5: "Exceptional"}


def _founder_insights_section(fi: models.FounderInsights) -> str:
    if not fi:
        return ""
    lines = ["## Founding Team Assessment\n"]
    if fi.ratings_json:
        try:
            ratings = json.loads(fi.ratings_json)
            for key, label in _RATING_LABELS.items():
                r = ratings.get(key, {})
                score = r.get("score")
                notes = r.get("notes", "").strip()
                if score or notes:
                    score_str = f"{_SCORE_LABELS.get(score, score)}" if score else ""
                    lines.append(f"**{label}**" + (f" — {score_str}" if score_str else ""))
                    if notes:
                        lines.append(notes)
                    lines.append("")
        except (json.JSONDecodeError, AttributeError):
            pass
    if fi.key_impressions and fi.key_impressions.strip():
        lines.append(f"**Key Impressions**\n{fi.key_impressions.strip()}\n")
    if fi.meeting_notes and fi.meeting_notes.strip():
        lines.append(f"**Meeting Notes**\n{fi.meeting_notes.strip()}\n")
    return "\n".join(lines) if len(lines) > 1 else ""


def _build_memo(deal: models.Deal, agent_outputs: list, feedback_entries: list, founder_insights: models.FounderInsights = None) -> str:
    by_agent = _latest_outputs_by_agent(agent_outputs)
    by_feedback = _latest_feedback_by_agent(feedback_entries)

    sd = by_agent.get("scientific_diligence", {})
    ci = by_agent.get("competitive_intelligence", {})
    cr = by_agent.get("clinical_regulatory", {})
    fv = by_agent.get("financing_valuation", {})

    feedback_block = ""
    for agent_label, key in [
        ("Scientific Diligence", "scientific_diligence"),
        ("Competitive Intelligence", "competitive_intelligence"),
        ("Clinical & Regulatory", "clinical_regulatory"),
        ("Financing & Valuation", "financing_valuation"),
    ]:
        if key in by_feedback:
            feedback_block += f"\n{agent_label} notes from analyst: {by_feedback[key]}"

    if founder_insights and founder_insights.ratings_json:
        try:
            ratings = json.loads(founder_insights.ratings_json)
            scored = [f"{_RATING_LABELS.get(k, k)}: {r.get('score')}/5" for k, r in ratings.items() if r.get("score")]
            if scored:
                feedback_block += f"\nFounding team assessment: {', '.join(scored)}"
        except (json.JSONDecodeError, AttributeError):
            pass
    if founder_insights and founder_insights.key_impressions:
        feedback_block += f"\nKey team impressions: {founder_insights.key_impressions}"

    agent_summary = json.dumps(
        {k: {ik: iv for ik, iv in v.items() if ik != "citations"} for k, v in by_agent.items()},
        indent=2,
    )

    prompt = f"""You are a senior biopharma investment analyst writing an investment memo.

Deal: {deal.company_name} | {deal.asset_name or "N/A"} | {deal.indication or "N/A"} | {deal.stage or "N/A"} | {deal.round_type or "N/A"}
Fund thesis: {deal.fund_thesis or "Not specified"}

Agent analysis:
{agent_summary}
{feedback_block}

Return ONLY valid JSON (no markdown fences):
{{
  "executive_summary": "3-4 sentence executive summary covering the opportunity, asset, and key investment thesis",
  "recommendation": "3-4 sentence recommendation covering the key reasons to proceed or not, with major caveats"
}}"""

    response = _client.messages.create(
        model=_SONNET,
        max_tokens=600,
        messages=[{"role": "user", "content": prompt}],
    )
    prose = _parse_json("".join(b.text for b in response.content if hasattr(b, "text") and b.type == "text"))

    exec_summary = prose.get("executive_summary") or deal.fund_thesis or "N/A"
    recommendation = prose.get("recommendation") or "N/A"

    all_questions = (
        sd.get("diligence_questions", [])
        + ci.get("diligence_questions", [])
        + cr.get("diligence_questions", [])
        + fv.get("diligence_questions", [])
    )

    all_citations = (
        sd.get("citations", []) + ci.get("citations", [])
        + cr.get("citations", []) + fv.get("citations", [])
    )
    source_notes = "\n".join(
        f"- [{c.get('filename', '?')} · chunk {c.get('chunk_index', '?')}] \"{c.get('quote', '')}\""
        for c in all_citations
    ) or "No documents indexed."

    analyst_notes_section = ""
    for label, key in [
        ("Scientific Diligence", "scientific_diligence"),
        ("Competitive Intelligence", "competitive_intelligence"),
        ("Clinical & Regulatory", "clinical_regulatory"),
        ("Financing & Valuation", "financing_valuation"),
    ]:
        if key in by_feedback:
            analyst_notes_section += f"\n**{label}**\n{by_feedback[key]}\n"

    fi_section = _founder_insights_section(founder_insights)

    memo = f"""# Investment Memo: {deal.company_name}

---

## Executive Summary

{exec_summary}

---

## Company Overview

| Field | Value |
|---|---|
| Company | {deal.company_name} |
| Asset | {deal.asset_name or "N/A"} |
| Indication | {deal.indication or "N/A"} |
| Stage | {deal.stage or "N/A"} |
| Round | {deal.round_type or "N/A"} |

---

## Scientific Diligence

**Mechanism of Action**
{sd.get("mechanism_of_action", "N/A")}

**Clinical Evidence**
{sd.get("clinical_evidence", "N/A")}

**Opportunities**
{_bullet(sd.get("scientific_opportunities", []))}

**Risks**
{_bullet(sd.get("scientific_risks", []))}

---

## Competitive Intelligence

**Market Overview**
{ci.get("market_overview", "N/A")}

**Differentiation**
{ci.get("differentiation", "N/A")}

**Opportunities**
{_bullet(ci.get("competitive_opportunities", []))}

**Risks**
{_bullet(ci.get("competitive_risks", []))}

---

## Clinical & Regulatory Assessment

**Regulatory Pathway**
{cr.get("regulatory_pathway", "N/A")}

**Precedent**
{cr.get("precedent", "N/A")}

**Opportunities**
{_bullet(cr.get("regulatory_opportunities", []))}

**Risks**
{_bullet(cr.get("regulatory_risks", []))}

---

## Financing & Valuation

**Comparable Financings**
{fv.get("comparable_financings", "N/A")}

**Valuation Considerations**
{fv.get("valuation_considerations", "N/A")}

**Opportunities**
{_bullet(fv.get("financing_opportunities", []))}

**Risks**
{_bullet(fv.get("financing_risks", []))}

---

## Consolidated Diligence Questions

{_bullet(all_questions)}

---

## Recommendation

{recommendation}
{f'''
---

{fi_section}''' if fi_section else ''}
{f'''
---

## Analyst Notes

{analyst_notes_section.strip()}''' if analyst_notes_section.strip() else ''}

---

## Source Notes

{source_notes}
"""
    return memo.strip()


@router.post("/{deal_id}/generate-memo", response_model=schemas.MemoOut)
def generate_memo(deal_id: int, db: Session = Depends(get_db)):
    deal = db.query(models.Deal).filter(models.Deal.id == deal_id).first()
    if not deal:
        raise HTTPException(status_code=404, detail="Deal not found")

    agent_outputs = (
        db.query(models.AgentOutput)
        .filter(models.AgentOutput.deal_id == deal_id)
        .order_by(models.AgentOutput.created_at.desc())
        .all()
    )
    if not agent_outputs:
        raise HTTPException(status_code=400, detail="Run agents first before generating a memo")

    feedback_entries = (
        db.query(models.AgentFeedback)
        .filter(models.AgentFeedback.deal_id == deal_id)
        .order_by(models.AgentFeedback.created_at.desc())
        .all()
    )

    fi = db.query(models.FounderInsights).filter(models.FounderInsights.deal_id == deal_id).first()

    markdown = _build_memo(deal, agent_outputs, feedback_entries, fi)
    memo = models.Memo(deal_id=deal_id, markdown=markdown)
    db.add(memo)
    db.commit()
    db.refresh(memo)
    return memo


def _add_runs(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def _markdown_to_docx(markdown: str, company_name: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:])
        elif line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not re.match(r"^\|[-| :]+\|$", lines[i].strip()):
                    table_lines.append(lines[i])
                i += 1
            if table_lines:
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                data_rows = [
                    [c.strip() for c in r.split("|")[1:-1]]
                    for r in table_lines[1:]
                ]
                tbl = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
                tbl.style = "Table Grid"
                for j, h in enumerate(headers):
                    tbl.rows[0].cells[j].text = h
                for ri, row in enumerate(data_rows):
                    for ci, val in enumerate(row):
                        if ci < len(tbl.rows[ri + 1].cells):
                            tbl.rows[ri + 1].cells[ci].text = val
                doc.add_paragraph()
            continue
        elif line == "---":
            doc.add_paragraph()
        elif line.strip():
            p = doc.add_paragraph()
            _add_runs(p, line)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@router.get("/{deal_id}/memo/export")
def export_memo_docx(deal_id: int, db: Session = Depends(get_db)):
    deal = db.query(models.Deal).filter(models.Deal.id == deal_id).first()
    if not deal:
        raise HTTPException(status_code=404, detail="Deal not found")
    memo = (
        db.query(models.Memo)
        .filter(models.Memo.deal_id == deal_id)
        .order_by(models.Memo.created_at.desc())
        .first()
    )
    if not memo:
        raise HTTPException(status_code=404, detail="No memo generated yet")

    docx_bytes = _markdown_to_docx(memo.markdown, deal.company_name)
    filename = f"{deal.company_name.replace(' ', '_')}_Investment_Memo.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{deal_id}/memo", response_model=schemas.MemoOut)
def get_memo(deal_id: int, db: Session = Depends(get_db)):
    memo = (
        db.query(models.Memo)
        .filter(models.Memo.deal_id == deal_id)
        .order_by(models.Memo.created_at.desc())
        .first()
    )
    if not memo:
        raise HTTPException(status_code=404, detail="No memo generated yet")
    return memo
